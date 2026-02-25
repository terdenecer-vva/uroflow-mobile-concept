#!/usr/bin/env python3
"""Extract and inventory Uroflow Project Package v4.2 ZIP contents."""

from __future__ import annotations

import argparse
import csv
import io
import re
import zipfile
from collections import defaultdict
from pathlib import Path

import openpyxl
from docx import Document

ZIP_ROOT_PREFIX = "Submission_Build_v4.2/"
AUTOMATION_PREFIX = f"{ZIP_ROOT_PREFIX}10_Pilot_Automation/"

SELECTED_XLSX = {
    "Uroflow_Master_Submission_Index_v4.2.xlsx",
    "Uroflow_Claims_Guardrails_Checklist_v1.0.xlsx",
    "Uroflow_QMS_Document_Register_v1.0.xlsx",
    "Uroflow_OnDevice_Privacy_Guardrails_Test_Cases_v1.0.xlsx",
    "Uroflow_Privacy_ContentGuardrails_v2_Test_Cases_v1.0.xlsx",
    "Uroflow_DHF_Freeze_Event_Log_Template_v1.0.xlsx",
    "Uroflow_Pilot_Freeze_Kit_Template_v1.0.xlsx",
    "Uroflow_Golden_Dataset_Coverage_Targets_Tracker_v1.0.xlsx",
    "Uroflow_Golden_Dataset_Weekly_Report_Template_v1.0.xlsx",
    "Ethics_Submission_Pack_Index_v1.0.xlsx",
    # Keep historical indexes for traceability across package lineage.
    "Uroflow_Master_Submission_Index_v3.8.xlsx",
    "Uroflow_Master_Submission_Index_v3.9.xlsx",
    "Uroflow_Master_Submission_Index_v4.0.xlsx",
    "Uroflow_Master_Submission_Index_v4.1.xlsx",
}

# Included for backward compatibility with earlier extraction script defaults.
LEGACY_SELECTED_XLSX = {
    "Uroflow_Launch_Readiness_Integrated_Plan_v1.0.xlsx",
    "Uroflow_Submission_Ready_EvidencePack_Index_EU_US_v1.0.xlsx",
    "Uroflow_Master_Submission_Index_v4.2.xlsx",
    "Uroflow_DHF_Index_Status_Register_v1.3_EXECUTED_AUTO.xlsx",
    "Uroflow_QMS_Forms_and_Registers_v1.1_EXECUTED.xlsx",
    "Uroflow_Acceptance_Test_Matrix_v1.0.xlsx",
    "Uroflow_P0_Pilot_Readiness_Evidence_Manifest_v1.0.xlsx",
    "Uroflow_Pilot_Issue_Log_v1.0.xlsx",
    "Uroflow_Golden_Dataset_Coverage_Targets_and_Tracker_v1.0.xlsx",
    "Uroflow_EU_MDR_GSPR_Checklist_AnnexI_v1.3_EXECUTED_EVIDENCE_AUTO.xlsx",
    "Uroflow_EU_MDR_AnnexII_III_Master_Index_v1.1_EXECUTED.xlsx",
    "EU_MDR_AnnexII_III_Submission_Folder_Index_v2.4.xlsx",
    "Uroflow_FDA_Software_Documentation_Index_v1.0.xlsx",
    "Uroflow_FDA_Software_Revision_History_Log_v1.1_BASELINE.xlsx",
    "Uroflow_FDA_Unresolved_Anomalies_Log_v1.1_BASELINE.xlsx",
    "Uroflow_Privacy_Compliance_Workbook_v1.0.xlsx",
    "Uroflow_Threat_Model_STRIDE_v1.0.xlsx",
    "Uroflow_HFE_Observation_Forms_and_Questionnaires_v1.0.xlsx",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--zip-path",
        type=Path,
        required=True,
        help="Path to Uroflow_Project_Package_v4.2.zip",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("docs/project-package-v4.2"),
        help="Output directory for inventories and extracted sheets.",
    )
    parser.add_argument(
        "--automation-out",
        type=Path,
        default=Path("scripts/pilot_automation_v4_2"),
        help="Output directory for extracted pilot-automation scripts/config.",
    )
    return parser.parse_args()


def file_category(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".xlsx"):
        return "xlsx"
    if lower.endswith(".csv"):
        return "csv"
    if lower.endswith(".zip"):
        return "zip"
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".txt"):
        return "text"
    if lower.endswith(".py"):
        return "python"
    if lower.endswith(".json"):
        return "json"
    if lower.endswith(".sh"):
        return "shell"
    if lower.endswith(".bat"):
        return "batch"
    if lower.endswith(".wav"):
        return "audio"
    return "other"


def module_name(path: str) -> str:
    if not path.startswith(ZIP_ROOT_PREFIX):
        return "unknown"
    remainder = path[len(ZIP_ROOT_PREFIX) :]
    parts = remainder.split("/")
    return parts[0] if parts else "unknown"


def clean_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_")


def write_csv(path: Path, fieldnames: list[str], rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def workbook_sheet_to_csv(
    workbook: openpyxl.Workbook,
    sheet_name: str,
    output_csv: Path,
) -> tuple[int, int]:
    worksheet = workbook[sheet_name]
    max_row = worksheet.max_row
    max_col = worksheet.max_column
    output_csv.parent.mkdir(parents=True, exist_ok=True)
    with output_csv.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        for row in worksheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
            writer.writerow(["" if cell.value is None else str(cell.value) for cell in row])
    return max_row, max_col


def docx_inventory_row(filename: str, file_bytes: bytes) -> dict[str, str]:
    document = Document(io.BytesIO(file_bytes))
    headings: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            headings.append(text)
    return {
        "file_name": filename,
        "paragraph_count": str(len(document.paragraphs)),
        "table_count": str(len(document.tables)),
        "heading_count": str(len(headings)),
        "heading_preview": " | ".join(headings[:12]),
    }


def should_extract_automation(member_name: str) -> bool:
    if not member_name.startswith(AUTOMATION_PREFIX):
        return False
    relative = member_name[len(AUTOMATION_PREFIX) :]
    if not relative:
        return False
    if "/__pycache__/" in f"/{relative}" or relative.endswith(".pyc"):
        return False
    if relative.startswith("outputs/"):
        return False
    if relative.startswith("sample/"):
        return True
    if relative.startswith(("scripts/", "config/", "schemas/")):
        return True
    base = Path(relative).name
    return (
        base.startswith("README_")
        or base.startswith("requirements")
        or base.startswith("run_")
    )


def main() -> None:
    args = parse_args()
    output_dir: Path = args.output_dir
    sheets_dir = output_dir / "sheets"
    automation_out: Path = args.automation_out
    output_dir.mkdir(parents=True, exist_ok=True)
    automation_out.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, str]] = []
    docx_rows: list[dict[str, str]] = []
    xlsx_rows: list[dict[str, str]] = []
    module_stats: dict[tuple[str, str], dict[str, int]] = defaultdict(
        lambda: {"file_count": 0, "size_bytes": 0}
    )
    extracted_automation_files = 0

    with zipfile.ZipFile(args.zip_path, "r") as archive:
        members = [item for item in archive.infolist() if not item.is_dir()]
        for info in sorted(members, key=lambda item: item.filename.lower()):
            name = info.filename
            category = file_category(name)
            module = module_name(name)
            is_selected_xlsx = (
                Path(name).name in SELECTED_XLSX
                or Path(name).name in LEGACY_SELECTED_XLSX
            ) and category == "xlsx"
            manifest_rows.append(
                {
                    "file_name": name,
                    "module": module,
                    "category": category,
                    "size_bytes": str(info.file_size),
                    "selected_sheet_export": "yes" if is_selected_xlsx else "no",
                    "automation_asset": "yes" if should_extract_automation(name) else "no",
                }
            )
            stats = module_stats[(module, category)]
            stats["file_count"] += 1
            stats["size_bytes"] += int(info.file_size)

            file_bytes = archive.read(name)
            if category == "docx":
                docx_rows.append(docx_inventory_row(name, file_bytes))

            if category == "xlsx":
                workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                for sheet_name in workbook.sheetnames:
                    xlsx_rows.append(
                        {
                            "file_name": name,
                            "sheet_name": sheet_name,
                            "sheet_rows": str(workbook[sheet_name].max_row),
                            "sheet_cols": str(workbook[sheet_name].max_column),
                            "sheet_exported": "yes" if is_selected_xlsx else "no",
                            "output_csv": "",
                        }
                    )

                if is_selected_xlsx:
                    path_key = clean_name(Path(name).with_suffix("").as_posix())
                    for sheet_name in workbook.sheetnames:
                        sheet_key = clean_name(sheet_name)
                        out_csv = sheets_dir / f"{path_key}__{sheet_key}.csv"
                        rows, cols = workbook_sheet_to_csv(workbook, sheet_name, out_csv)
                        xlsx_rows.append(
                            {
                                "file_name": name,
                                "sheet_name": sheet_name,
                                "sheet_rows": str(rows),
                                "sheet_cols": str(cols),
                                "sheet_exported": "yes",
                                "output_csv": str(out_csv.relative_to(output_dir)),
                            }
                        )

            if should_extract_automation(name):
                relative = name[len(AUTOMATION_PREFIX) :]
                destination = automation_out / relative
                destination.parent.mkdir(parents=True, exist_ok=True)
                destination.write_bytes(file_bytes)
                extracted_automation_files += 1

        readme_path = f"{ZIP_ROOT_PREFIX}00_README_and_Indexes/00_README_Submission_Build_v4.2.txt"
        if readme_path in archive.namelist():
            readme_text = archive.read(readme_path).decode("utf-8", errors="replace")
            (output_dir / "source_readme_v4.2.txt").write_text(readme_text, encoding="utf-8")

    module_rows: list[dict[str, str]] = []
    for (module, category), stats in sorted(module_stats.items()):
        module_rows.append(
            {
                "module": module,
                "category": category,
                "file_count": str(stats["file_count"]),
                "size_bytes": str(stats["size_bytes"]),
            }
        )

    write_csv(
        output_dir / "package_manifest_v4.2.csv",
        [
            "file_name",
            "module",
            "category",
            "size_bytes",
            "selected_sheet_export",
            "automation_asset",
        ],
        manifest_rows,
    )
    write_csv(
        output_dir / "module_summary_v4.2.csv",
        ["module", "category", "file_count", "size_bytes"],
        module_rows,
    )
    write_csv(
        output_dir / "docx_inventory_v4.2.csv",
        ["file_name", "paragraph_count", "table_count", "heading_count", "heading_preview"],
        docx_rows,
    )
    write_csv(
        output_dir / "xlsx_sheet_inventory_v4.2.csv",
        ["file_name", "sheet_name", "sheet_rows", "sheet_cols", "sheet_exported", "output_csv"],
        xlsx_rows,
    )

    print(f"Inventory exported to: {output_dir}")
    print(f"Manifest rows: {len(manifest_rows)}")
    print(f"DOCX inventory rows: {len(docx_rows)}")
    print(f"XLSX inventory rows: {len(xlsx_rows)}")
    print(f"Automation assets extracted: {extracted_automation_files}")
    print(f"Automation output dir: {automation_out}")


if __name__ == "__main__":
    main()
