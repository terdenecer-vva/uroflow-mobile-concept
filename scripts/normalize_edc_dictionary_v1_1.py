#!/usr/bin/env python3
"""Build canonical EDC v1.1 artifacts from ZIP + DOCX + repo eCRF sources."""

from __future__ import annotations

import argparse
import csv
import io
import re
import zipfile
from collections import defaultdict
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency for local generation
    Document = None


FIELD_ALIASES = {
    "sex": "sex_at_birth",
    "posture": "voiding_position",
    "primary_condition": "diagnostic_group",
    "ambient_noise_level": "environment_noise_level",
    "lighting": "lighting_level",
    "session_datetime_local": "visit_datetime",
    "qmax_ref": "ref_qmax_ml_s",
    "qavg_ref": "ref_qavg_ml_s",
    "vvoid_ref": "ref_vvoid_ml",
    "flowtime_ref": "ref_flow_time_s",
    "tqmax_ref": "ref_tqmax_s",
    "pattern_ref": "ref_curve_class",
    "qmax_app": "app_qmax_ml_s",
    "qavg_app": "app_qavg_ml_s",
    "vvoid_app": "app_vvoid_ml",
    "flowtime_app": "app_flow_time_s",
    "tqmax_app": "app_tqmax_s",
    "pattern_app": "app_curve_class",
    "quality_reasons": "quality_flags",
    "invalid_flag": "quality_status",
    "invalid_reason": "quality_notes",
    "uncertainty_overall": "overall_uncertainty",
    "pvr_done": "pvr_available",
    "pvr_minutes_after_void": "pvr_delay_min",
    "qmax": "app_qmax_ml_s",
    "qavg": "app_qavg_ml_s",
    "vvoid": "app_vvoid_ml",
    "flow_time": "app_flow_time_s",
    "tqmax": "app_tqmax_s",
    "pattern_class": "app_curve_class",
    "pvr_measured": "pvr_available",
    "toilet_scan_performed": "toilet_scan_done",
    "pour_calibration_used": "pour_calibration_done",
    "mode": "capture_mode",
    "audio_sample_rate": "audio_sample_rate_hz",
    "recording_duration": "recording_duration_s",
    "flow_start_time": "flow_start_time_s",
    "flow_end_time": "flow_end_time_s",
    "q_curve": "q_curve_series",
    "q_curve_uncertainty": "q_curve_uncertainty_series",
}

SECTION_OVERRIDES = {
    "study_id": "identification",
    "site_id": "identification",
    "subject_id": "identification",
    "visit_id": "identification",
    "session_id": "identification",
    "visit_datetime": "identification",
    "timezone": "identification",
    "operator_id": "identification",
    "toilet_id": "capture_setup",
    "device_model": "capture_setup",
    "ios_version": "capture_setup",
    "app_version": "capture_setup",
    "model_id": "capture_setup",
    "model_version": "capture_setup",
    "model_hash": "capture_setup",
    "lidar_available": "capture_setup",
    "capture_fps": "capture_setup",
    "audio_sample_rate_hz": "capture_setup",
    "toilet_scan_done": "capture_setup",
    "toilet_fingerprint_id": "capture_setup",
    "pour_calibration_done": "capture_setup",
    "pour_calibration_volume_ml": "capture_setup",
    "recording_duration_s": "capture_setup",
    "capture_mode": "capture_setup",
    "age_years": "screening",
    "sex_at_birth": "screening",
    "height_cm": "screening",
    "weight_kg": "screening",
    "diagnostic_group": "screening",
    "medications": "screening",
    "consent_signed": "screening",
    "consent_level": "screening",
    "voiding_position": "test_conditions",
    "void_into_water": "test_conditions",
    "environment_noise_level": "test_conditions",
    "lighting_level": "test_conditions",
    "noise_source": "test_conditions",
    "phone_on_support": "test_conditions",
    "flush_during_recording": "test_conditions",
    "attempt_number": "test_conditions",
    "comments_conditions": "test_conditions",
    "flow_start_time_s": "event_timing",
    "flow_end_time_s": "event_timing",
    "ref_qmax_ml_s": "reference_uroflow",
    "ref_qavg_ml_s": "reference_uroflow",
    "ref_vvoid_ml": "reference_uroflow",
    "ref_flow_time_s": "reference_uroflow",
    "ref_tqmax_s": "reference_uroflow",
    "ref_curve_class": "reference_uroflow",
    "ref_artifacts": "reference_uroflow",
    "app_qmax_ml_s": "app_output",
    "app_qavg_ml_s": "app_output",
    "app_vvoid_ml": "app_output",
    "app_flow_time_s": "app_output",
    "app_tqmax_s": "app_output",
    "app_curve_class": "app_output",
    "q_curve_sampling_hz": "app_output",
    "q_curve_series": "app_output",
    "q_curve_uncertainty_series": "app_output",
    "intermittency_index": "app_output",
    "quality_score": "quality",
    "qr_low_snr": "quality",
    "qr_motion": "quality",
    "qr_roi_lost": "quality",
    "qr_not_in_water": "quality",
    "qr_low_volume": "quality",
    "qr_other_text": "quality",
    "quality_flags": "quality",
    "quality_status": "quality",
    "quality_notes": "quality",
    "overall_uncertainty": "quality",
    "motion_score": "quality",
    "roi_visibility": "quality",
    "flush_detected": "quality",
    "repeat_required": "quality",
    "repeat_reason": "quality",
    "protocol_deviation": "quality",
    "deviation_comment": "quality",
    "pvr_available": "pvr",
    "pvr_method": "pvr",
    "pvr_delay_min": "pvr",
    "pvr_ml": "pvr",
    "media_saved": "privacy_export",
    "media_consent": "privacy_export",
    "export_format": "privacy_export",
    "export_destination": "privacy_export",
    "notes_clinician": "privacy_export",
    "ae_occurred": "safety",
    "ae_description": "safety",
    "ae_related": "safety",
    "delta_qmax": "derived",
    "delta_qavg": "derived",
    "delta_vvoid": "derived",
    "abs_pct_error_qmax": "derived",
}

TYPE_OVERRIDES = {
    "age_years": "int",
    "attempt_number": "int",
    "capture_fps": "int",
    "audio_sample_rate_hz": "int",
    "quality_score": "int",
    "q_curve_series": "float[]",
    "q_curve_uncertainty_series": "float[]",
    "quality_flags": "enum[]",
    "quality_status": "enum",
    "consent_signed": "bool",
    "lidar_available": "bool",
    "toilet_scan_done": "bool",
    "pour_calibration_done": "bool",
    "phone_on_support": "bool",
    "flush_during_recording": "bool",
    "pvr_available": "bool",
    "media_saved": "bool",
    "media_consent": "bool",
    "ae_occurred": "bool",
    "ae_related": "bool",
    "repeat_required": "bool",
    "protocol_deviation": "bool",
    "flush_detected": "bool",
    "delta_qmax": "float",
    "delta_qavg": "float",
    "delta_vvoid": "float",
    "abs_pct_error_qmax": "float",
}

UNIT_OVERRIDES = {
    "app_qmax_ml_s": "ml/s",
    "app_qavg_ml_s": "ml/s",
    "ref_qmax_ml_s": "ml/s",
    "ref_qavg_ml_s": "ml/s",
    "app_vvoid_ml": "ml",
    "ref_vvoid_ml": "ml",
    "pvr_ml": "ml",
    "pvr_delay_min": "min",
    "app_flow_time_s": "s",
    "app_tqmax_s": "s",
    "ref_flow_time_s": "s",
    "ref_tqmax_s": "s",
    "flow_start_time_s": "s",
    "flow_end_time_s": "s",
    "recording_duration_s": "s",
    "q_curve_sampling_hz": "Hz",
    "audio_sample_rate_hz": "Hz",
    "capture_fps": "fps",
    "quality_score": "score",
    "overall_uncertainty": "0-1",
}

CODELIST_OVERRIDES = {
    "sex_at_birth": "SEX_AT_BIRTH",
    "voiding_position": "VOIDING_POSITION",
    "diagnostic_group": "DIAGNOSTIC_GROUP",
    "lidar_available": "YES_NO",
    "toilet_scan_done": "YES_NO",
    "pour_calibration_done": "YES_NO",
    "consent_signed": "YES_NO",
    "void_into_water": "VOID_INTO_WATER",
    "environment_noise_level": "NOISE_LEVEL",
    "lighting_level": "LIGHTING_LEVEL",
    "phone_on_support": "YES_NO",
    "flush_during_recording": "YES_NO",
    "ref_curve_class": "FLOW_PATTERN",
    "app_curve_class": "FLOW_PATTERN",
    "quality_status": "QUALITY_STATUS",
    "quality_flags": "QUALITY_FLAGS",
    "roi_visibility": "ROI_VISIBILITY",
    "pvr_available": "YES_NO",
    "pvr_method": "PVR_METHOD",
    "ae_occurred": "YES_NO",
    "ae_related": "YES_NO",
    "capture_mode": "CAPTURE_MODE",
    "export_format": "EXPORT_FORMAT",
    "export_destination": "EXPORT_DESTINATION",
    "consent_level": "CONSENT_LEVEL",
}

DERIVED_FIELDS = {"delta_qmax", "delta_qavg", "delta_vvoid", "abs_pct_error_qmax"}

ZIP_LIST_NAME_ALIASES = {
    "YESNO": "YES_NO",
    "SEX": "SEX_AT_BIRTH",
    "POSTURE": "VOIDING_POSITION",
    "PRIMARY_CONDITION": "DIAGNOSTIC_GROUP",
    "NOISE_LEVEL": "NOISE_LEVEL",
    "LIGHTING": "LIGHTING_LEVEL",
    "VOID_INTO_WATER": "VOID_INTO_WATER",
    "PATTERN": "FLOW_PATTERN",
    "QUALITY_REASONS": "QUALITY_FLAGS",
    "PVR_METHOD": "PVR_METHOD",
}

ZIP_CODE_ALIASES = {
    ("YES_NO", "1"): "yes",
    ("YES_NO", "0"): "no",
    ("SEX_AT_BIRTH", "M"): "male",
    ("SEX_AT_BIRTH", "F"): "female",
    ("VOIDING_POSITION", "M_STAND"): "male_standing",
    ("VOIDING_POSITION", "M_SIT"): "male_sitting",
    ("VOIDING_POSITION", "F_SIT"): "female_sitting",
    ("NOISE_LEVEL", "NONE"): "none",
    ("NOISE_LEVEL", "MODERATE"): "moderate",
    ("NOISE_LEVEL", "HIGH"): "high",
    ("LIGHTING_LEVEL", "NORMAL"): "normal",
    ("LIGHTING_LEVEL", "DIM"): "dim",
    ("LIGHTING_LEVEL", "GLARE"): "glare",
    ("VOID_INTO_WATER", "YES"): "yes",
    ("VOID_INTO_WATER", "NO"): "no",
    ("VOID_INTO_WATER", "UNSURE"): "unsure",
    ("FLOW_PATTERN", "BELL"): "bell",
    ("FLOW_PATTERN", "PLATEAU"): "plateau",
    ("FLOW_PATTERN", "INTERMITTENT"): "intermittent",
    ("FLOW_PATTERN", "OTHER"): "other",
    ("QUALITY_FLAGS", "NOISE"): "noise",
    ("QUALITY_FLAGS", "MOTION"): "motion",
    ("QUALITY_FLAGS", "ROI_LOST"): "roi_lost",
    ("QUALITY_FLAGS", "NOT_IN_WATER"): "not_in_water",
    ("QUALITY_FLAGS", "FLUSH"): "flush",
    ("QUALITY_FLAGS", "SHORT_VOID"): "low_volume",
    ("QUALITY_FLAGS", "OTHER"): "other",
    ("PVR_METHOD", "BLADDER_SCAN"): "bladder_scan",
    ("PVR_METHOD", "ULTRASOUND"): "ultrasound",
    ("DIAGNOSTIC_GROUP", "BPH_LUTS"): "bph_luts",
    ("DIAGNOSTIC_GROUP", "NEURO"): "neurogenic",
    ("DIAGNOSTIC_GROUP", "STRICTURE"): "stricture",
    ("DIAGNOSTIC_GROUP", "POSTOP"): "postop",
    ("DIAGNOSTIC_GROUP", "OTHER"): "other",
}

ADDITIONAL_CODELIST_ROWS = [
    ("QUALITY_STATUS", "valid", "Valid", "Валидно", "repo"),
    ("QUALITY_STATUS", "repeat", "Repeat needed", "Нужен повтор", "repo"),
    ("QUALITY_STATUS", "reject", "Rejected", "Отклонено", "repo"),
    ("CAPTURE_MODE", "water_impact", "Water impact", "В воду", "docx"),
    ("CAPTURE_MODE", "jet_assist", "Jet in air assist", "Струя в воздухе", "docx"),
    ("CAPTURE_MODE", "fallback_nonwater", "Fallback non-water", "Не в воду", "docx"),
    ("ROI_VISIBILITY", "ok", "ROI visible", "ROI виден", "docx"),
    ("ROI_VISIBILITY", "partial", "ROI partially visible", "ROI частично виден", "docx"),
    ("ROI_VISIBILITY", "lost", "ROI lost", "ROI потерян", "docx"),
    ("CONSENT_LEVEL", "metrics_only", "Metrics only", "Только метрики", "docx"),
    ("CONSENT_LEVEL", "roi_video_local", "ROI video local", "Локальное ROI-видео", "docx"),
    ("CONSENT_LEVEL", "full_recording", "Full recording", "Полная запись", "docx"),
    ("EXPORT_FORMAT", "pdf", "PDF report", "PDF-отчёт", "docx"),
    ("EXPORT_FORMAT", "fhir", "FHIR Observation", "FHIR", "docx"),
    ("EXPORT_FORMAT", "hl7", "HL7 ORU^R01", "HL7", "docx"),
    ("EXPORT_DESTINATION", "local", "Local device", "Локально", "docx"),
    ("EXPORT_DESTINATION", "emr", "EMR/EHR", "EMR/EHR", "docx"),
]

TRANSFORM_NOTES = {
    ("zip", "invalid_flag"): "map YES_NO to QUALITY_STATUS: yes->reject, no->valid",
    ("zip", "uncertainty_overall"): "rename to overall_uncertainty",
    ("zip", "quality_reasons"): "rename to quality_flags",
    ("zip", "pvr_done"): "rename to pvr_available",
    ("zip", "pvr_minutes_after_void"): "rename to pvr_delay_min",
    ("docx", "qmax"): "rename to app_qmax_ml_s",
    ("docx", "qavg"): "rename to app_qavg_ml_s",
    ("docx", "vvoid"): "rename to app_vvoid_ml",
    ("docx", "flow_time"): "rename to app_flow_time_s",
    ("docx", "tqmax"): "rename to app_tqmax_s",
    ("docx", "pattern_class"): "rename to app_curve_class",
    ("docx", "session_datetime_local"): "rename to visit_datetime",
    ("docx", "audio_sample_rate"): "rename to audio_sample_rate_hz",
    ("docx", "q_curve"): "rename to q_curve_series",
    ("docx", "q_curve_uncertainty"): "rename to q_curve_uncertainty_series",
}

VALIDATION_RULES = [
    (
        "VR-001",
        "error",
        "app_qmax_ml_s >= app_qavg_ml_s",
        "Qmax приложения должен быть не меньше Qavg.",
        "clinical baseline",
    ),
    (
        "VR-002",
        "error",
        "ref_qmax_ml_s >= ref_qavg_ml_s",
        "Эталонный Qmax должен быть не меньше эталонного Qavg.",
        "clinical baseline",
    ),
    (
        "VR-003",
        "error",
        "pvr_available == no => pvr_ml is null and pvr_method is null",
        "При отсутствии PVR поля PVR должны быть пустыми.",
        "repo v0.1",
    ),
    (
        "VR-004",
        "error",
        "pvr_available == yes => pvr_method is not null",
        "При наличии PVR метод обязателен.",
        "clinical protocol",
    ),
    (
        "VR-005",
        "error",
        "quality_status in (repeat,reject) => repeat_reason is not null",
        "Для repeat/reject требуется причина повтора.",
        "repo v0.1",
    ),
    (
        "VR-006",
        "error",
        "protocol_deviation == yes => deviation_comment is not null",
        "При отклонении от протокола нужен комментарий.",
        "repo v0.1",
    ),
    (
        "VR-007",
        "warning",
        "quality_score < 60 => quality_status != valid",
        "Низкий quality_score несовместим со статусом valid.",
        "quality policy",
    ),
    (
        "VR-008",
        "error",
        "flow_end_time_s > flow_start_time_s",
        "Время окончания потока должно быть больше времени старта.",
        "docx v1.0",
    ),
    (
        "VR-009",
        "warning",
        "pour_calibration_done == yes => pour_calibration_volume_ml is not null",
        "При pour calibration должен быть указан объём.",
        "capture SOP",
    ),
    (
        "VR-010",
        "warning",
        "void_into_water == no => capture_mode != water_impact",
        "Режим water_impact некорректен без попадания в воду.",
        "capture SOP",
    ),
    (
        "VR-011",
        "warning",
        "app_vvoid_ml >= 50",
        "Малый объём мочеиспускания может быть нерепрезентативным.",
        "clinical QC",
    ),
    (
        "VR-012",
        "warning",
        "capture_fps >= 24 and audio_sample_rate_hz in (44100, 48000)",
        "Скорость видео/аудио ниже минимальных значений.",
        "FR-010",
    ),
]


@dataclass
class SourceField:
    source: str
    field_name: str
    section: str
    type_name: str
    units: str
    required: str
    codelist: str
    description: str
    phi_sensitive: str

    @property
    def canonical_field(self) -> str:
        return FIELD_ALIASES.get(self.field_name, self.field_name)


@dataclass
class CanonicalField:
    name: str
    source_fields: list[SourceField] = field(default_factory=list)

    def add(self, source_field: SourceField) -> None:
        self.source_fields.append(source_field)

    def required_level(self) -> str:
        levels = {normalize_required(x.required) for x in self.source_fields}
        if "M" in levels:
            return "M"
        if "C" in levels:
            return "C"
        return "O"

    def source_coverage(self) -> str:
        return ",".join(sorted({x.source for x in self.source_fields}))

    def aliases(self) -> str:
        names = sorted({x.field_name for x in self.source_fields if x.field_name != self.name})
        return ",".join(names)

    def section(self) -> str:
        if self.name in SECTION_OVERRIDES:
            return SECTION_OVERRIDES[self.name]
        for field_item in self.source_fields:
            if field_item.section:
                return slugify(field_item.section)
        return "unclassified"

    def type_name(self) -> str:
        if self.name in TYPE_OVERRIDES:
            return TYPE_OVERRIDES[self.name]
        types_by_source = defaultdict(list)
        for field_item in self.source_fields:
            normalized = normalize_type(field_item.type_name, field_item.source)
            if normalized:
                types_by_source[field_item.source].append(normalized)
        for preferred_source in ("repo", "zip", "docx"):
            if types_by_source.get(preferred_source):
                return types_by_source[preferred_source][0]
        return "string"

    def units(self) -> str:
        if self.name in UNIT_OVERRIDES:
            return UNIT_OVERRIDES[self.name]
        for preferred_source in ("zip", "docx", "repo"):
            for field_item in self.source_fields:
                if field_item.source == preferred_source and field_item.units:
                    return field_item.units
        return ""

    def codelist(self) -> str:
        if self.name in CODELIST_OVERRIDES:
            return CODELIST_OVERRIDES[self.name]
        for preferred_source in ("zip", "docx", "repo"):
            for field_item in self.source_fields:
                if field_item.source == preferred_source and field_item.codelist:
                    return canonical_list_name(field_item.codelist)
        return ""

    def phi_sensitive(self) -> str:
        values = {normalize_phi(x.phi_sensitive) for x in self.source_fields}
        return "yes" if "yes" in values else "no"

    def description_ru(self) -> str:
        for preferred_source in ("zip", "docx", "repo"):
            for field_item in self.source_fields:
                if field_item.source == preferred_source and field_item.description:
                    return normalize_spaces(field_item.description)
        return ""


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9_]+", "_", text.lower()).strip("_")


def normalize_required(value: str) -> str:
    normalized = normalize_spaces(value).lower()
    if normalized in {"m", "y", "yes", "true", "да"}:
        return "M"
    if normalized in {"c", "conditional"}:
        return "C"
    return "O"


def normalize_phi(value: str) -> str:
    normalized = normalize_spaces(value).lower()
    if normalized in {"yes", "y", "true", "1", "да"}:
        return "yes"
    return "no"


def normalize_type(value: str, source: str) -> str:
    normalized = normalize_spaces(value).lower()
    if not normalized:
        return ""
    zip_map = {
        "text": "string",
        "number": "float",
        "datetime": "datetime",
        "dropdown": "enum",
        "multiselect": "enum[]",
    }
    docx_map = {
        "string": "string",
        "number": "float",
        "integer": "int",
        "bool": "bool",
        "enum": "enum",
        "array<number>": "float[]",
        "array<enum>": "enum[]",
        "datetime": "datetime",
    }
    repo_map = {
        "string": "string",
        "int": "int",
        "float": "float",
        "bool": "bool",
        "enum": "enum",
        "datetime": "datetime",
    }
    if source == "zip":
        return zip_map.get(normalized, normalized)
    if source == "docx":
        return docx_map.get(normalized, normalized)
    return repo_map.get(normalized, normalized)


def canonical_list_name(source_list_name: str) -> str:
    if not source_list_name:
        return ""
    return ZIP_LIST_NAME_ALIASES.get(source_list_name, source_list_name.upper())


def parse_zip_source(zip_path: Path) -> tuple[list[SourceField], list[dict[str, str]]]:
    source_fields: list[SourceField] = []
    codelists: list[dict[str, str]] = []
    with zipfile.ZipFile(zip_path) as archive:
        data_dictionary_name = next(
            name for name in archive.namelist() if name.lower().endswith("/datadictionary.csv")
        )
        codelists_name = next(
            name for name in archive.namelist() if name.lower().endswith("/codelists.csv")
        )
        with archive.open(data_dictionary_name) as file_pointer:
            reader = csv.DictReader(io.TextIOWrapper(file_pointer, encoding="utf-8-sig"))
            for row in reader:
                source_fields.append(
                    SourceField(
                        source="zip",
                        field_name=row["FieldName"].strip(),
                        section=row["Form"].strip(),
                        type_name=row["Type"].strip(),
                        units=row["Units"].strip(),
                        required=row["Required"].strip(),
                        codelist=row["Choices/List"].strip(),
                        description=row["Label_RU"].strip(),
                        phi_sensitive=row["PHI/Sensitive"].strip(),
                    )
                )
        with archive.open(codelists_name) as file_pointer:
            reader = csv.DictReader(io.TextIOWrapper(file_pointer, encoding="utf-8-sig"))
            codelists = [dict(row) for row in reader]
    return source_fields, codelists


def parse_docx_source(docx_path: Path | None) -> list[SourceField]:
    if docx_path is None:
        return []
    if Document is None:
        raise RuntimeError(
            "python-docx is required for --docx-path. "
            "Install with pip install python-docx."
        )
    document = Document(str(docx_path))
    if not document.tables:
        return []
    headers = [normalize_spaces(cell.text) for cell in document.tables[0].rows[0].cells]
    result: list[SourceField] = []
    for row in document.tables[0].rows[1:]:
        cells = [normalize_spaces(cell.text) for cell in row.cells]
        data = dict(zip(headers, cells, strict=False))
        field_name = data.get("Field", "")
        if not field_name:
            continue
        result.append(
            SourceField(
                source="docx",
                field_name=field_name,
                section="docx_package",
                type_name=data.get("Тип/формат", ""),
                units=data.get("Ед.", ""),
                required=data.get("Req.", ""),
                codelist="",
                description=data.get("Описание", ""),
                phi_sensitive=data.get("PHI/PII", ""),
            )
        )
    return result


def parse_repo_markdown(repo_md_path: Path) -> list[SourceField]:
    fields: list[SourceField] = []
    current_section = "repo_ecrf"
    for raw_line in repo_md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if line.startswith("## "):
            current_section = line[3:].strip()
            continue
        if not (line.startswith("|") and "`" in line):
            continue
        parts = [item.strip() for item in line.strip("|").split("|")]
        if not parts:
            continue
        field_cell = parts[0]
        field_match = re.match(r"`([^`]+)`", field_cell)
        if not field_match:
            continue
        field_name = field_match.group(1).strip()
        if len(parts) >= 4:
            type_name = parts[1].strip()
            required = parts[2].strip()
            description = parts[3].strip()
        else:
            type_name = "float" if field_name in DERIVED_FIELDS else "string"
            required = "O"
            description = parts[2].strip() if len(parts) >= 3 else ""
        fields.append(
            SourceField(
                source="repo",
                field_name=field_name,
                section=current_section,
                type_name=type_name,
                units="",
                required=required,
                codelist="",
                description=description,
                phi_sensitive="",
            )
        )
    return fields


def build_catalog(all_fields: Iterable[SourceField]) -> dict[str, CanonicalField]:
    catalog: dict[str, CanonicalField] = {}
    for source_field in all_fields:
        canonical = source_field.canonical_field
        if canonical not in catalog:
            catalog[canonical] = CanonicalField(name=canonical)
        catalog[canonical].add(source_field)
    return catalog


def build_codelists_rows(zip_codelists: list[dict[str, str]]) -> list[dict[str, str]]:
    rows = []
    seen = set()
    for row in zip_codelists:
        canonical_list = canonical_list_name(row["ListName"])
        canonical_code = ZIP_CODE_ALIASES.get((canonical_list, row["Code"]), row["Code"].lower())
        out_row = {
            "list_name": canonical_list,
            "canonical_code": canonical_code,
            "label_en": canonical_code.replace("_", " ").title(),
            "label_ru": row["Label_RU"],
            "source_system": "zip",
            "source_list_name": row["ListName"],
            "source_code": row["Code"],
        }
        key = (
            out_row["list_name"],
            out_row["canonical_code"],
            out_row["source_system"],
            out_row["source_code"],
        )
        if key not in seen:
            seen.add(key)
            rows.append(out_row)

    for list_name, code, label_en, label_ru, source_system in ADDITIONAL_CODELIST_ROWS:
        out_row = {
            "list_name": list_name,
            "canonical_code": code,
            "label_en": label_en,
            "label_ru": label_ru,
            "source_system": source_system,
            "source_list_name": list_name,
            "source_code": code,
        }
        key = (
            out_row["list_name"],
            out_row["canonical_code"],
            out_row["source_system"],
            out_row["source_code"],
        )
        if key not in seen:
            seen.add(key)
            rows.append(out_row)
    rows.sort(key=lambda row: (row["list_name"], row["canonical_code"], row["source_system"]))
    return rows


def build_crosswalk_rows(catalog: dict[str, CanonicalField]) -> list[dict[str, str]]:
    rows = []
    for canonical_name in sorted(catalog):
        field_entry = catalog[canonical_name]
        for source_field in sorted(
            field_entry.source_fields, key=lambda item: (item.source, item.field_name)
        ):
            mapping_type = "direct"
            if source_field.field_name in DERIVED_FIELDS:
                mapping_type = "derived"
            elif source_field.field_name != canonical_name:
                mapping_type = "rename"
            if (source_field.source, source_field.field_name) in TRANSFORM_NOTES:
                mapping_type = "mapped"
            rows.append(
                {
                    "canonical_field": canonical_name,
                    "source_system": source_field.source,
                    "source_field": source_field.field_name,
                    "mapping_type": mapping_type,
                    "transform_note": TRANSFORM_NOTES.get(
                        (source_field.source, source_field.field_name), ""
                    ),
                }
            )
    return rows


def build_dictionary_rows(catalog: dict[str, CanonicalField]) -> list[dict[str, str]]:
    rows = []
    for canonical_name in sorted(catalog):
        entry = catalog[canonical_name]
        rows.append(
            {
                "field_name": canonical_name,
                "section": entry.section(),
                "type": entry.type_name(),
                "units": entry.units(),
                "required_level": entry.required_level(),
                "codelist": entry.codelist(),
                "phi_sensitive": entry.phi_sensitive(),
                "description_ru": entry.description_ru(),
                "source_coverage": entry.source_coverage(),
                "aliases": entry.aliases(),
            }
        )
    return rows


def build_gap_rows(catalog: dict[str, CanonicalField]) -> list[dict[str, str]]:
    rows = []
    for canonical_name in sorted(catalog):
        sources = {source_field.source for source_field in catalog[canonical_name].source_fields}
        rows.append(
            {
                "field_name": canonical_name,
                "has_zip": "yes" if "zip" in sources else "no",
                "has_docx": "yes" if "docx" in sources else "no",
                "has_repo": "yes" if "repo" in sources else "no",
                "coverage_count": str(len(sources)),
            }
        )
    return rows


def write_csv(path: Path, fieldnames: list[str], rows: Iterable[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def write_readme(
    path: Path, catalog: dict[str, CanonicalField], zip_count: int, docx_count: int, repo_count: int
) -> None:
    coverage_distribution = defaultdict(int)
    for item in catalog.values():
        coverage_distribution[len({x.source for x in item.source_fields})] += 1
    lines = [
        "# EDC v1.1 Normalization Package",
        "",
        "Generated by `scripts/normalize_edc_dictionary_v1_1.py`.",
        "",
        "## Inputs",
        f"- ZIP DataDictionary fields: {zip_count}",
        f"- DOCX package fields: {docx_count}",
        f"- Repo eCRF fields: {repo_count}",
        f"- Canonical fields (union after alias mapping): {len(catalog)}",
        "",
        "## Coverage",
        f"- Fields present in 1 source: {coverage_distribution[1]}",
        f"- Fields present in 2 sources: {coverage_distribution[2]}",
        f"- Fields present in all 3 sources: {coverage_distribution[3]}",
        "",
        "## Files",
        "- `canonical_data_dictionary_v1.1.csv`: normalized field-level schema.",
        "- `field_crosswalk_v1.1.csv`: source-to-canonical mapping with transform notes.",
        "- `canonical_codelists_v1.1.csv`: merged value sets and source aliases.",
        "- `validation_rules_v1.1.csv`: cross-field and quality validation rules.",
        "- `source_coverage_v1.1.csv`: source coverage matrix per canonical field.",
        "",
        "## Regeneration",
        "```bash",
        "python3 scripts/normalize_edc_dictionary_v1_1.py \\",
        "  --zip-path '/path/to/Uroflow_EDC_CRF_DataDictionary_v1.0_CSV.zip' \\",
        "  --docx-path '/path/to/Uroflow_Smartphone_Design_Clinical_Package_v1.0.docx' \\",
        "  --repo-md docs/ethics-package-v0.1/ecrf-data-dictionary-v0.1.en.md \\",
        "  --output-dir docs/edc-v1.1",
        "```",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--zip-path",
        type=Path,
        required=True,
        help="Path to EDC DataDictionary ZIP",
    )
    parser.add_argument("--docx-path", type=Path, help="Path to clinical package DOCX")
    parser.add_argument(
        "--repo-md",
        type=Path,
        default=Path("docs/ethics-package-v0.1/ecrf-data-dictionary-v0.1.en.md"),
        help="Path to existing repo markdown dictionary",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("docs/edc-v1.1"),
        help="Target folder for canonical artifacts",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    zip_fields, zip_codelists = parse_zip_source(args.zip_path)
    docx_fields = parse_docx_source(args.docx_path)
    repo_fields = parse_repo_markdown(args.repo_md)

    all_fields = [*zip_fields, *docx_fields, *repo_fields]
    catalog = build_catalog(all_fields)

    dictionary_rows = build_dictionary_rows(catalog)
    crosswalk_rows = build_crosswalk_rows(catalog)
    codelist_rows = build_codelists_rows(zip_codelists)
    gap_rows = build_gap_rows(catalog)

    rules_rows = [
        {
            "rule_id": rule_id,
            "severity": severity,
            "expression": expression,
            "error_message_ru": message,
            "source": source,
        }
        for rule_id, severity, expression, message, source in VALIDATION_RULES
    ]

    output_dir = args.output_dir
    write_csv(
        output_dir / "canonical_data_dictionary_v1.1.csv",
        [
            "field_name",
            "section",
            "type",
            "units",
            "required_level",
            "codelist",
            "phi_sensitive",
            "description_ru",
            "source_coverage",
            "aliases",
        ],
        dictionary_rows,
    )
    write_csv(
        output_dir / "field_crosswalk_v1.1.csv",
        ["canonical_field", "source_system", "source_field", "mapping_type", "transform_note"],
        crosswalk_rows,
    )
    write_csv(
        output_dir / "canonical_codelists_v1.1.csv",
        [
            "list_name",
            "canonical_code",
            "label_en",
            "label_ru",
            "source_system",
            "source_list_name",
            "source_code",
        ],
        codelist_rows,
    )
    write_csv(
        output_dir / "validation_rules_v1.1.csv",
        ["rule_id", "severity", "expression", "error_message_ru", "source"],
        rules_rows,
    )
    write_csv(
        output_dir / "source_coverage_v1.1.csv",
        ["field_name", "has_zip", "has_docx", "has_repo", "coverage_count"],
        gap_rows,
    )
    write_readme(
        output_dir / "README.md",
        catalog,
        zip_count=len(zip_fields),
        docx_count=len(docx_fields),
        repo_count=len(repo_fields),
    )

    print(f"Generated EDC v1.1 package in: {output_dir}")
    print(f"Canonical fields: {len(catalog)}")


if __name__ == "__main__":
    main()
