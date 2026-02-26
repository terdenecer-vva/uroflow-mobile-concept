#!/usr/bin/env python3
"""Extract and inventory Uroflow Project Package v1.1 ZIP contents."""

from __future__ import annotations

import argparse
import csv
import io
import re
import zipfile
from pathlib import Path

import openpyxl
from docx import Document

SELECTED_XLSX = {
    "Uroflow_Bench_BOM_and_Wiring_v1.0.xlsx",
    "Uroflow_Bench_and_Dataset_Templates_v1.0.xlsx",
    "Uroflow_Regulatory_Mapping_Matrix_v1.0.xlsx",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--zip-path",
        type=Path,
        required=True,
        help="Path to Uroflow_Project_Package_v1.1.zip",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("docs/project-package-v1.1"),
        help="Output directory for extracted artifacts",
    )
    return parser.parse_args()


def clean_sheet_name(sheet_name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", sheet_name).strip("_")


def file_category(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".xlsx"):
        return "xlsx"
    if lower.endswith(".csv.zip") or lower.endswith(".zip"):
        return "zip"
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".txt"):
        return "text"
    return "other"


def write_csv(path: Path, fieldnames: list[str], rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def workbook_sheet_to_csv(workbook, sheet_name: str, out_path: Path) -> tuple[int, int]:
    worksheet = workbook[sheet_name]
    max_row, max_col = worksheet.max_row, worksheet.max_column
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        for row in worksheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
            writer.writerow(["" if cell.value is None else str(cell.value) for cell in row])
    return max_row, max_col


def extract_docx_inventory(filename: str, file_bytes: bytes) -> dict[str, str]:
    document = Document(io.BytesIO(file_bytes))
    headings = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            headings.append(text)
    heading_preview = " | ".join(headings[:12])
    return {
        "file_name": filename,
        "paragraph_count": str(len(document.paragraphs)),
        "table_count": str(len(document.tables)),
        "heading_preview": heading_preview,
    }


def main() -> None:
    args = parse_args()
    output_dir = args.output_dir
    sheets_dir = output_dir / "sheets"
    output_dir.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, str]] = []
    docx_rows: list[dict[str, str]] = []
    xlsx_rows: list[dict[str, str]] = []

    with zipfile.ZipFile(args.zip_path, "r") as archive:
        members = [info for info in archive.infolist() if not info.is_dir()]
        for info in sorted(members, key=lambda item: item.filename.lower()):
            name = info.filename
            category = file_category(name)
            manifest_rows.append(
                {
                    "file_name": name,
                    "size_bytes": str(info.file_size),
                    "category": category,
                    "selected_for_extraction": "yes" if Path(name).name in SELECTED_XLSX else "no",
                }
            )

            if category == "docx":
                file_bytes = archive.read(name)
                docx_rows.append(extract_docx_inventory(Path(name).name, file_bytes))

            if category == "xlsx" and Path(name).name in SELECTED_XLSX:
                file_bytes = archive.read(name)
                workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
                stem = Path(name).stem
                for sheet_name in workbook.sheetnames:
                    out_csv = sheets_dir / f"{stem}__{clean_sheet_name(sheet_name)}.csv"
                    row_count, col_count = workbook_sheet_to_csv(workbook, sheet_name, out_csv)
                    xlsx_rows.append(
                        {
                            "file_name": Path(name).name,
                            "sheet_name": sheet_name,
                            "rows": str(row_count),
                            "cols": str(col_count),
                            "output_csv": str(out_csv.relative_to(output_dir)),
                        }
                    )

        if "00_README_Uroflow_Project_Package_v1.1.txt" in archive.namelist():
            readme_text = archive.read("00_README_Uroflow_Project_Package_v1.1.txt").decode(
                "utf-8", errors="replace"
            )
            (output_dir / "source_readme_v1.1.txt").write_text(readme_text, encoding="utf-8")

    write_csv(
        output_dir / "package_manifest_v1.1.csv",
        ["file_name", "size_bytes", "category", "selected_for_extraction"],
        manifest_rows,
    )
    write_csv(
        output_dir / "docx_inventory_v1.1.csv",
        ["file_name", "paragraph_count", "table_count", "heading_preview"],
        docx_rows,
    )
    write_csv(
        output_dir / "xlsx_sheet_inventory_v1.1.csv",
        ["file_name", "sheet_name", "rows", "cols", "output_csv"],
        xlsx_rows,
    )

    print(f"Inventory exported to: {output_dir}")
    print(f"Manifest files: {len(manifest_rows)}")
    print(f"DOCX inventory rows: {len(docx_rows)}")
    print(f"XLSX sheet exports: {len(xlsx_rows)}")


if __name__ == "__main__":
    main()
