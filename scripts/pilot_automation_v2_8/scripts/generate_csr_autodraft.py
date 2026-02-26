#!/usr/bin/env python3
"""
CSR DOCX auto-draft generator (offline)

Purpose
- Take TFL outputs produced by run_tfl_from_golden_dataset.py:
    * tfl_summary.json
    * tfl_record_level.csv
    * ba_plots/BA_<Metric>.png (optional)
- Insert numbered Tables / Figures / Listings into a CSR template DOCX that contains
  placeholder tokens.

Default placeholders (CSR template v1.1_AUTOFILL*)
- {{AUTO_TABLE_T1_QMAX}}
- {{AUTO_FIGURE_F1_QMAX}}
- {{AUTO_TABLE_T2_QAVG}}
- {{AUTO_FIGURE_F2_QAVG}}
- {{AUTO_TABLE_T3_VVOID}}
- {{AUTO_FIGURE_F3_VVOID}}
- {{AUTO_TABLE_T4_FLOWTIME}}
- {{AUTO_FIGURE_F4_FLOWTIME}}
- {{AUTO_LISTING_L1}}

Config
- By default reads ../config/csr_autodraft_config.json (relative to this script).
  You can override with --config.

Output
- DOCX with embedded tables and images
- Meta JSON (for audit trail)

Example
python scripts/generate_csr_autodraft.py \
  --tfl_dir outputs/tfl \
  --csr_template ../../05_Clinical/Uroflow_CSR_Template_v1.1_AUTOFILL.docx \
  --out_dir outputs/csr_autodraft --lang EN
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


# ---------------------------
# Helpers
# ---------------------------

def load_json(p: Path) -> dict:
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)

def _is_nan(x) -> bool:
    try:
        return x != x
    except Exception:
        return False

def _fmt(x: Optional[float], nd: int = 3) -> str:
    if x is None or _is_nan(x):
        return ""
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x)

def _fmt_pct(x: Optional[float], nd: int = 2) -> str:
    if x is None or _is_nan(x):
        return ""
    try:
        return f"{float(x):.{nd}f}%"
    except Exception:
        return str(x)

def _move_after(new_element, ref_element):
    ref_element.addnext(new_element)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

def insert_table_after_paragraph(doc: Document, paragraph, rows: List[List[str]]):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for r in rows:
        row_cells = table.add_row().cells
        for j, val in enumerate(r):
            row_cells[j].text = str(val)
            for run in row_cells[j].paragraphs[0].runs:
                run.font.size = Pt(9)
    # bold header row
    for j in range(len(rows[0])):
        for run in table.rows[0].cells[j].paragraphs[0].runs:
            run.font.bold = True
    _move_after(table._tbl, paragraph._p)
    return table

def insert_paragraph_after_element(doc: Document, ref_element, text: str, italic: bool = False, align_center: bool = False):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].italic = italic
        p.runs[0].font.size = Pt(9)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _move_after(p._p, ref_element)
    return p

def insert_picture_after_element(doc: Document, ref_element, img_path: Path, width_in: float = 5.8):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _move_after(p._p, ref_element)
    return p


# ---------------------------
# Content builders
# ---------------------------

def build_agreement_table(stats: dict, lang: str, metric_label: str, unit: str) -> List[List[str]]:
    if lang.upper() == "RU":
        header = ["Показатель", "Ед.", "n", "Смещение (bias)", "SD", "LoA низ", "LoA верх", "MAE", "MAPE"]
    else:
        header = ["Metric", "Unit", "n", "Bias", "SD", "LoA low", "LoA high", "MAE", "MAPE"]
    row = [
        metric_label,
        unit,
        str(stats.get("n","")),
        _fmt(stats.get("bias")),
        _fmt(stats.get("sd")),
        _fmt(stats.get("loa_low")),
        _fmt(stats.get("loa_high")),
        _fmt(stats.get("mae")),
        _fmt_pct(stats.get("mape")),
    ]
    return [header, row]

def build_listing_excerpt(df, preferred_cols: List[str], max_rows: int = 20) -> List[List[str]]:
    cols = [c for c in preferred_cols if c in df.columns]
    if not cols:
        cols = list(df.columns)[:10]
    dfx = df[cols].head(max_rows).copy()
    # format numeric columns
    for c in dfx.columns:
        if getattr(dfx[c].dtype, "kind", "") in ("i","f"):
            dfx[c] = dfx[c].map(lambda x: "" if pd.isna(x) else round(float(x), 3))
    header = cols
    rows = [header] + dfx.astype(str).values.tolist()
    return rows


# ---------------------------
# Main
# ---------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--tfl_dir", required=True, help="Folder containing tfl_summary.json and tfl_record_level.csv")
    ap.add_argument("--csr_template", required=True, help="CSR DOCX template with placeholder tokens")
    ap.add_argument("--out_dir", default="outputs/csr_autodraft", help="Output directory")
    ap.add_argument("--lang", default="EN", choices=["EN","RU"], help="Language for inserted tables/notes")
    ap.add_argument("--config", default=str(Path(__file__).resolve().parents[1] / "config" / "csr_autodraft_config.json"))
    ap.add_argument("--plots_subdir", default="", help="Override plots subdir if needed")
    ap.add_argument("--max_listing_rows", type=int, default=20)
    args = ap.parse_args()

    if pd is None:
        raise SystemExit("pandas is required for listing formatting")

    tfl_dir = Path(args.tfl_dir)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    summary_path = tfl_dir / "tfl_summary.json"
    listing_path = tfl_dir / "tfl_record_level.csv"
    if not summary_path.exists():
        raise SystemExit(f"Missing: {summary_path}")
    if not listing_path.exists():
        raise SystemExit(f"Missing: {listing_path}")

    cfg = {}
    cfg_path = Path(args.config)
    if cfg_path.exists():
        cfg = load_json(cfg_path)

    plots_subdir = args.plots_subdir or cfg.get("plots_subdir", "ba_plots")
    plots_dir = tfl_dir / plots_subdir

    placeholder_cfg = cfg.get("placeholders", {})
    if not placeholder_cfg:
        # fallback to defaults (same as template tokens)
        placeholder_cfg = {
            "{{AUTO_TABLE_T1_QMAX}}": {"kind":"table","metric":"Qmax","unit":"ml/s","label_en":"Qmax","label_ru":"Qmax"},
            "{{AUTO_FIGURE_F1_QMAX}}": {"kind":"fig","metric":"Qmax","filename":"BA_Qmax.png"},
            "{{AUTO_TABLE_T2_QAVG}}": {"kind":"table","metric":"Qavg","unit":"ml/s","label_en":"Qavg","label_ru":"Qavg"},
            "{{AUTO_FIGURE_F2_QAVG}}": {"kind":"fig","metric":"Qavg","filename":"BA_Qavg.png"},
            "{{AUTO_TABLE_T3_VVOID}}": {"kind":"table","metric":"Vvoid","unit":"ml","label_en":"Vvoid","label_ru":"Vvoid"},
            "{{AUTO_FIGURE_F3_VVOID}}": {"kind":"fig","metric":"Vvoid","filename":"BA_Vvoid.png"},
            "{{AUTO_TABLE_T4_FLOWTIME}}": {"kind":"table","metric":"FlowTime","unit":"s","label_en":"Flow time","label_ru":"Время потока"},
            "{{AUTO_FIGURE_F4_FLOWTIME}}": {"kind":"fig","metric":"FlowTime","filename":"BA_FlowTime.png"},
            "{{AUTO_LISTING_L1}}": {"kind":"listing","max_rows":20},
        }

    preferred_cols = cfg.get("listing_columns_preferred", ["record_id","site_id","toilet_id","quality_score","valid_for_primary"])
    summary = load_json(summary_path)
    df_listing = pd.read_csv(listing_path)

    doc = Document(args.csr_template)

    # Iterate through paragraphs. For each placeholder paragraph, insert content AFTER it, then delete it.
    for p in list(doc.paragraphs):
        token = p.text.strip()
        if token not in placeholder_cfg:
            continue
        item = placeholder_cfg[token]
        kind = str(item.get("kind","")).lower()

        if kind == "table":
            metric = str(item.get("metric",""))
            unit = str(item.get("unit",""))
            label = item.get("label_ru", metric) if args.lang.upper() == "RU" else item.get("label_en", metric)
            stats = summary.get("metrics", {}).get(metric, {})
            rows = build_agreement_table(stats, args.lang, label, unit)
            table = insert_table_after_paragraph(doc, p, rows)
            # optional note
            if stats.get("n", 0) in (0, 1, 2):
                note = "Insufficient N for agreement stats." if args.lang.upper() == "EN" else "Недостаточно наблюдений для расчёта показателей согласия."
                insert_paragraph_after_element(doc, table._tbl, note, italic=True)

        elif kind == "fig":
            fname = str(item.get("filename",""))
            img_path = plots_dir / fname
            if img_path.exists():
                pic_p = insert_picture_after_element(doc, p._p, img_path)
            else:
                note = "Plot not available (run TFL with --make_plots)." if args.lang.upper() == "EN" else "График отсутствует (запустите TFL с --make_plots)."
                insert_paragraph_after_element(doc, p._p, note, italic=True)

        elif kind == "listing":
            max_rows = int(item.get("max_rows", args.max_listing_rows))
            rows = build_listing_excerpt(df_listing, preferred_cols, max_rows=max_rows)
            table = insert_table_after_paragraph(doc, p, rows)
            note = f"Full record-level listing: {listing_path.name}" if args.lang.upper() == "EN" else f"Полный листинг: {listing_path.name}"
            insert_paragraph_after_element(doc, table._tbl, note, italic=True)

        else:
            note = f"Unknown placeholder kind: {kind}"
            insert_paragraph_after_element(doc, p._p, note, italic=True)

        # Remove placeholder paragraph
        delete_paragraph(p)

    out_name = f"Uroflow_CSR_Autodraft_{args.lang}_{tfl_dir.name}.docx"
    out_path = out_dir / out_name
    doc.save(out_path)

    meta = {
        "version": "1.0",
        "lang": args.lang,
        "tfl_dir": str(tfl_dir),
        "csr_template": str(Path(args.csr_template).name),
        "output_docx": str(out_path),
        "n_total": summary.get("n_total"),
        "n_valid": summary.get("n_valid"),
        "plots_dir": str(plots_dir),
        "config": str(cfg_path),
    }
    with open(out_dir / f"csr_autodraft_meta_{args.lang}.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    print(f"OK: {out_path}")

if __name__ == "__main__":
    main()
