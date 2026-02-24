#!/usr/bin/env python3
"""
G1 evidence bundle builder (offline)

Runs (or consumes) golden dataset performance outputs and produces a G1 bundle:
- Runs TFL generator (record listing + BA/MAE + optional plots)
- Runs drift dashboard
- Evaluates acceptance criteria (config-driven)
- Produces:
  - G1_evidence_summary.xlsx
  - Pilot-freeze V&V report (executed) in RU and EN (auto-filled evidence table)

This script is designed to help make Gate G1 a repeatable, auditable process.

Usage (example)
python build_g1_evidence_bundle.py \
  --dataset_root <DATASET_ROOT> --manifest <MANIFEST.csv> \
  --submission_build_root <SUBMISSION_BUILD_ROOT> \
  --out_dir <OUT>

If submission_build_root is provided, the script will locate:
- CSR TFL template: 05_Clinical/Uroflow_CSR_TFL_Workbook_v1.0.xlsx
- Pilot-freeze V&V templates: 01_Product_QMS/Uroflow_Software_System_VV_Report_Pilot_Freeze_v1.0_RU.docx and EN.docx
"""

from __future__ import annotations

import argparse
import json
import math
import shutil
from pathlib import Path
from typing import Dict

import pandas as pd

from run_tfl_from_golden_dataset import ba_stats
from run_tfl_from_golden_dataset import main as tfl_main  # for reuse
import run_tfl_from_golden_dataset as tfl_mod
import run_drift_dashboard as drift_mod


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(obj: dict, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(obj, indent=2, ensure_ascii=False), encoding="utf-8")


def evaluate_g1(df: pd.DataFrame, g1_cfg: dict) -> Dict[str, dict]:
    crit = g1_cfg["criteria"]
    out = {}

    n_total = int(df.shape[0])
    n_valid = int((df["valid_for_primary"] == "Y").sum())
    valid_rate = (n_valid / n_total) if n_total else math.nan

    dv = df[df["valid_for_primary"] == "Y"]
    st_qmax = ba_stats(dv["ref_Qmax_ml_s"].to_numpy(float), dv["app_Qmax_ml_s"].to_numpy(float)) if dv.shape[0] >= 3 else {"mae": math.nan, "mape": math.nan, "n": n_valid}
    st_qavg = ba_stats(dv["ref_Qavg_ml_s"].to_numpy(float), dv["app_Qavg_ml_s"].to_numpy(float)) if dv.shape[0] >= 3 else {"mae": math.nan, "mape": math.nan, "n": n_valid}
    st_v = ba_stats(dv["ref_Vvoid_ml"].to_numpy(float), dv["app_Vvoid_ml"].to_numpy(float)) if dv.shape[0] >= 3 else {"mae": math.nan, "mape": math.nan, "n": n_valid}
    st_ft = ba_stats(dv["ref_FlowTime_s"].to_numpy(float), dv["app_FlowTime_s"].to_numpy(float)) if dv.shape[0] >= 3 else {"mae": math.nan, "mape": math.nan, "n": n_valid}

    def _pass(value, thresh, mode="le"):
        if not math.isfinite(value):
            return False
        return value <= thresh if mode == "le" else value >= thresh

    out["valid_rate"] = {"value": valid_rate, "threshold": crit["valid_rate_min"], "pass": _pass(valid_rate, crit["valid_rate_min"], mode="ge")}
    out["mae_qmax"] = {"value": st_qmax.get("mae", math.nan), "threshold": crit["mae_qmax_max_ml_s"], "pass": _pass(st_qmax.get("mae", math.nan), crit["mae_qmax_max_ml_s"])}
    out["mae_qavg"] = {"value": st_qavg.get("mae", math.nan), "threshold": crit["mae_qavg_max_ml_s"], "pass": _pass(st_qavg.get("mae", math.nan), crit["mae_qavg_max_ml_s"])}
    out["mape_vvoid"] = {"value": st_v.get("mape", math.nan), "threshold": crit["mape_vvoid_max"], "pass": _pass(st_v.get("mape", math.nan), crit["mape_vvoid_max"])}
    out["mae_flowtime"] = {"value": st_ft.get("mae", math.nan), "threshold": crit["mae_flowtime_max_s"], "pass": _pass(st_ft.get("mae", math.nan), crit["mae_flowtime_max_s"])}

    out["_counts"] = {"n_total": n_total, "n_valid": n_valid}
    return out


def write_g1_summary_xlsx(g1_eval: dict, out_path: Path, artifacts: dict) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "G1_Summary"
    ws.append(["Criterion", "Value", "Threshold", "Pass"])
    for k, v in g1_eval.items():
        if k.startswith("_"):
            continue
        ws.append([k, v["value"], v["threshold"], "PASS" if v["pass"] else "FAIL"])

    ws2 = wb.create_sheet("Counts")
    ws2.append(["n_total", g1_eval["_counts"]["n_total"]])
    ws2.append(["n_valid", g1_eval["_counts"]["n_valid"]])

    ws3 = wb.create_sheet("Artifacts")
    ws3.append(["Artifact", "Path"])
    for k, p in artifacts.items():
        ws3.append([k, str(p)])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def fill_vv_report(template_path: Path, out_path: Path, g1_eval: dict, artifacts: dict) -> None:
    from docx import Document

    doc = Document(template_path)
    doc.add_paragraph("")
    doc.add_paragraph("AUTO-GENERATED EVIDENCE TABLE (G1)")
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Evidence item"
    hdr[1].text = "Metric"
    hdr[2].text = "Value"
    hdr[3].text = "Threshold"
    hdr[4].text = "Result"

    for k, v in g1_eval.items():
        if k.startswith("_"):
            continue
        row = table.add_row().cells
        row[0].text = "G1 acceptance check"
        row[1].text = k
        row[2].text = f"{v['value']}"
        row[3].text = f"{v['threshold']}"
        row[4].text = "PASS" if v["pass"] else "FAIL"

    doc.add_paragraph("")
    doc.add_paragraph("AUTO-GENERATED ARTIFACT LINKS")
    for name, p in artifacts.items():
        doc.add_paragraph(f"- {name}: {p}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--dataset_root", required=True)
    ap.add_argument("--manifest", required=True)
    ap.add_argument("--submission_build_root", default="", help="Optional: root of Submission_Build_vX.Y folder")
    ap.add_argument("--out_dir", default=str(Path(__file__).resolve().parents[1] / "outputs" / "g1"))
    ap.add_argument("--g1_config", default=str(Path(__file__).resolve().parents[1] / "config" / "g1_acceptance_config.json"))
    ap.add_argument("--make_plots", action="store_true")
    ap.add_argument("--make_pdf", action="store_true")
    args = ap.parse_args()

    out_dir = Path(args.out_dir).expanduser().resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    g1_cfg = load_json(Path(args.g1_config))

    # Run TFL
    tfl_out_dir = out_dir / "tfl"
    tfl_out_dir.mkdir(parents=True, exist_ok=True)
    # call module main logic by constructing argv-like args
    # Use CSR template if submission_build_root provided
    csr_template = ""
    if args.submission_build_root:
        sb = Path(args.submission_build_root).expanduser().resolve()
        cand = sb / "05_Clinical" / "Uroflow_CSR_TFL_Workbook_v1.0.xlsx"
        if cand.exists():
            csr_template = str(cand)

    # Run tfl via subprocess-free direct call: set up args and call functions
    # We'll call the script's main by temporarily patching sys.argv
    import sys
    old_argv = sys.argv[:]
    try:
        sys.argv = ["run_tfl_from_golden_dataset.py",
                    "--dataset_root", args.dataset_root,
                    "--manifest", args.manifest,
                    "--out_dir", str(tfl_out_dir)]
        if csr_template:
            sys.argv += ["--csr_template", csr_template, "--csr_out", str(tfl_out_dir / "Uroflow_CSR_TFL_Workbook_filled.xlsx")]
        if args.make_plots:
            sys.argv += ["--make_plots"]
        if args.make_pdf:
            sys.argv += ["--make_pdf"]
        tfl_mod.main()
    finally:
        sys.argv = old_argv

    tfl_csv = tfl_out_dir / "tfl_record_level.csv"
    if not tfl_csv.exists():
        raise SystemExit("TFL record-level CSV not generated")

    df = pd.read_csv(tfl_csv)
    g1_eval = evaluate_g1(df, g1_cfg)

    # Drift dashboard
    drift_out_dir = out_dir / "drift"
    drift_out_dir.mkdir(parents=True, exist_ok=True)

    import sys
    old_argv = sys.argv[:]
    try:
        sys.argv = ["run_drift_dashboard.py",
                    "--tfl_csv", str(tfl_csv),
                    "--out_dir", str(drift_out_dir),
                    "--g1_config", str(Path(args.g1_config))]
        drift_mod.main()
    finally:
        sys.argv = old_argv

    artifacts = {
        "tfl_record_level_csv": tfl_csv,
        "tfl_summary_json": tfl_out_dir / "tfl_summary.json",
        "csr_filled_workbook": tfl_out_dir / "Uroflow_CSR_TFL_Workbook_filled.xlsx",
        "drift_dashboard_xlsx": drift_out_dir / "drift_dashboard.xlsx",
    }

    # Write G1 summary workbook
    g1_xlsx = out_dir / "G1_Evidence_Summary.xlsx"
    write_g1_summary_xlsx(g1_eval, g1_xlsx, artifacts)

    # Fill V&V report templates if submission root provided
    if args.submission_build_root:
        sb = Path(args.submission_build_root).expanduser().resolve()
        ru_tpl = sb / "01_Product_QMS" / "Uroflow_Software_System_VV_Report_Pilot_Freeze_v1.0_RU.docx"
        en_tpl = sb / "01_Product_QMS" / "Uroflow_Software_System_VV_Report_Pilot_Freeze_v1.0_EN.docx"
        if ru_tpl.exists():
            fill_vv_report(ru_tpl, out_dir / "Uroflow_SVV_Report_PilotFreeze_EXECUTED_RU.docx", g1_eval, artifacts)
        if en_tpl.exists():
            fill_vv_report(en_tpl, out_dir / "Uroflow_SVV_Report_PilotFreeze_EXECUTED_EN.docx", g1_eval, artifacts)

    # Save eval JSON
    save_json(g1_eval, out_dir / "g1_eval.json")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
